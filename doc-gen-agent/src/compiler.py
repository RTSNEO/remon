import os
import glob
from docx import Document
from docx.shared import Pt
import re

def compile_markdown_to_docx(input_dir: str, output_file: str):
    """Compiles a directory of markdown files into a single docx file."""
    print(f"Compiling markdown files from {input_dir} into {output_file}...")
    document = Document()
    
    # Optional styling
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Sort files to ensure correct order
    md_files = sorted(glob.glob(os.path.join(input_dir, "*.md")))
    
    if not md_files:
        print(f"No markdown files found in {input_dir}")
        return
        
    for file_path in md_files:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            
            # Simple markdown parsing for docx integration
            # Note: For production use with tables etc., you might want to use python-markdown 
            # or pandoc instead, but this covers basic headings and text structure.
            lines = content.splitlines()
            for line in lines:
                if line.startswith("# "):
                    document.add_heading(line[2:].strip(), level=1)
                elif line.startswith("## "):
                    document.add_heading(line[3:].strip(), level=2)
                elif line.startswith("### "):
                    document.add_heading(line[4:].strip(), level=3)
                elif line.startswith("- "):
                    document.add_paragraph(line[2:], style='List Bullet')
                elif line.strip() == "":
                    continue # Skip empty lines
                else:
                    document.add_paragraph(line)
        
        # document.add_page_break() # Optional: add page breaks between major sections
        
    document.save(output_file)
    print(f"Saved compiled document to {output_file}")


if __name__ == "__main__":
    # Example usage
    # compile_markdown_to_docx("../output/srd_drafts", "../output/Final_SRD.docx")
    pass
